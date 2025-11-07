from pathlib import Path
from typing import Optional, Dict
import io
import uuid
from datetime import datetime

from pypdf import PdfReader
from docx import Document as DocxDocument
from sqlalchemy.orm import Session

from models.document import Document
from core.config import settings

class IngestionService:
    def __init__(self, db: Session):
        self.db = db
        self.upload_dir = settings.UPLOAD_DIR_PATH

    async def ingest_document(
        self,
        file_content: bytes,
        original_filename: str
    ) -> Document:
        file_ext = Path(original_filename).suffix.lower()
        if file_ext not in settings.ALLOWED_EXTENSIONS:
            raise ValueError(f"File type not allowed. Allowed: {settings.ALLOWED_EXTENSIONS}")

        file_size = len(file_content)
        if file_size > settings.MAX_FILE_SIZE:
            raise ValueError(f"File too large. Maximum: {settings.MAX_FILE_SIZE} bytes")

        unique_filename = f"{uuid.uuid4()}{file_ext}"
        file_path = self.upload_dir / unique_filename

        with open(file_path, 'wb') as f:
            f.write(file_content)

        file_type = file_ext.lstrip('.')

        try:
            if file_type == 'pdf':
                content, metadata = self._extract_pdf(file_content)
            elif file_type == 'docx':
                content, metadata = self._extract_docx(file_content)
            elif file_type == 'txt':
                content, metadata = self._extract_txt(file_content)
            elif file_type == 'md':
                content, metadata = self._extract_markdown(file_content)
            else:
                raise ValueError(f"Tipo de arquivo não suportado: {file_type}")

            preview = content[:500] if len(content) > 500 else content

            document = Document(
                filename=unique_filename,
                original_filename=original_filename,
                file_type=file_type,
                file_size=file_size,
                file_path=str(file_path),
                content=content,
                content_preview=preview,
                num_pages=metadata.get('num_pages'),
                num_words=self._count_words(content),
                num_characters=len(content),
                language=self._detect_language(content),
                is_processed=False,
                processing_status='uploaded'
            )

            self.db.add(document)
            self.db.commit()
            self.db.refresh(document)

            return document

        except Exception as e:
            if file_path.exists():
                file_path.unlink()
            raise ValueError(f"Erro ao processar arquivo: {str(e)}")

    def ingest_document_sync(
        self,
        file_content: bytes,
        original_filename: str
    ) -> Document:
        file_ext = Path(original_filename).suffix.lower()
        if file_ext not in settings.ALLOWED_EXTENSIONS:
            raise ValueError(f"File type not allowed. Allowed: {settings.ALLOWED_EXTENSIONS}")

        file_size = len(file_content)
        if file_size > settings.MAX_FILE_SIZE:
            raise ValueError(f"File too large. Maximum: {settings.MAX_FILE_SIZE} bytes")

        unique_filename = f"{uuid.uuid4()}{file_ext}"
        file_path = self.upload_dir / unique_filename

        with open(file_path, 'wb') as f:
            f.write(file_content)

        file_type = file_ext.lstrip('.')

        try:
            if file_type == 'pdf':
                content, metadata = self._extract_pdf(file_content)
            elif file_type == 'docx':
                content, metadata = self._extract_docx(file_content)
            elif file_type == 'txt':
                content, metadata = self._extract_txt(file_content)
            elif file_type == 'md':
                content, metadata = self._extract_markdown(file_content)
            else:
                raise ValueError(f"Tipo de arquivo não suportado: {file_type}")

            preview = content[:500] if len(content) > 500 else content

            document = Document(
                filename=unique_filename,
                original_filename=original_filename,
                file_type=file_type,
                file_size=file_size,
                file_path=str(file_path),
                content=content,
                content_preview=preview,
                num_pages=metadata.get('num_pages'),
                num_words=self._count_words(content),
                num_characters=len(content),
                language=self._detect_language(content),
                is_processed=False,
                processing_status='uploaded'
            )

            self.db.add(document)
            self.db.commit()
            self.db.refresh(document)

            return document

        except Exception as e:
            if file_path.exists():
                file_path.unlink()
            raise ValueError(f"Erro ao processar arquivo: {str(e)}")

    def _extract_pdf(self, file_content: bytes) -> tuple[str, Dict]:

        pdf_file = io.BytesIO(file_content)
        reader = PdfReader(pdf_file)

        text_parts = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)

        content = "\n\n".join(text_parts)

        metadata = {
            'num_pages': len(reader.pages)
        }

        return content, metadata

    def _extract_docx(self, file_content: bytes) -> tuple[str, Dict]:

        docx_file = io.BytesIO(file_content)
        doc = DocxDocument(docx_file)

        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        content = "\n\n".join(text_parts)

        metadata = {
            'num_pages': None
        }

        return content, metadata

    def _extract_txt(self, file_content: bytes) -> tuple[str, Dict]:

        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        content = None

        for encoding in encodings:
            try:
                content = file_content.decode(encoding)
                break
            except UnicodeDecodeError:
                continue

        if content is None:
            raise ValueError("Não foi possível decodificar o arquivo TXT")

        metadata = {
            'num_pages': None
        }

        return content, metadata

    def _extract_markdown(self, file_content: bytes) -> tuple[str, Dict]:

        return self._extract_txt(file_content)

    def _count_words(self, text: str) -> int:

        return len(text.split())

    def _detect_language(self, text: str) -> Optional[str]:

        pt_words = ['o', 'a', 'de', 'que', 'e', 'do', 'da', 'em', 'um', 'para']
        en_words = ['the', 'be', 'to', 'of', 'and', 'a', 'in', 'that', 'have', 'it']

        text_lower = text.lower()[:1000]
        words = text_lower.split()

        pt_count = sum(1 for word in words if word in pt_words)
        en_count = sum(1 for word in words if word in en_words)

        if pt_count > en_count:
            return 'pt'
        elif en_count > pt_count:
            return 'en'
        else:
            return None

    def get_document(self, document_id: int) -> Optional[Document]:

        return self.db.query(Document).filter(Document.id == document_id).first()

    def list_documents(self, skip: int = 0, limit: int = 100) -> list[Document]:

        return self.db.query(Document).offset(skip).limit(limit).all()

    def delete_document(self, document_id: int) -> bool:

        document = self.get_document(document_id)
        if not document:
            return False

        file_path = Path(str(document.file_path))
        if file_path.exists():
            file_path.unlink()

        self.db.delete(document)
        self.db.commit()

        return True

    def get_stats(self) -> Dict:

        total = self.db.query(Document).count()

        total_size = self.db.query(
            func.sum(Document.file_size)
        ).scalar() or 0

        by_type = {}
        for file_type in ['pdf', 'docx', 'txt', 'md']:
            count = self.db.query(Document).filter(
                Document.file_type == file_type
            ).count()
            by_type[file_type] = count

        by_status = {}
        for status in ['uploaded', 'processing', 'completed', 'failed']:
            count = self.db.query(Document).filter(
                Document.processing_status == status
            ).count()
            by_status[status] = count

        return {
            'total_documents': total,
            'total_size_mb': round(total_size / (1024 * 1024), 2),
            'by_type': by_type,
            'by_status': by_status
        }

from sqlalchemy import func
